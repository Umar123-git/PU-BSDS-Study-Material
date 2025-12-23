import os
import traceback
import tempfile
import csv
import datetime
import io
import json
from pathlib import Path
import gradio as gr


from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain_core.documents import Document as LC_Document

try:
    import docx
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# Try to use ftfy to repair mojibake; fall back to replacing U+FFFD
try:
    import ftfy
    def _fix_text(s: str):
        if s is None:
            return s
        out = ftfy.fix_text(s)
        try:
            import re
            out = re.sub(r'(?<=[A-Za-z])\ufffd(?=[A-Za-z])', 't', out)
        except Exception:
            pass
        return out
except Exception:
    def _fix_text(s: str):
        if s is None:
            return s
        out = s.replace('\ufffd', '?')
        try:
            import re
            out = re.sub(r'(?<=[A-Za-z])\?+(?=[A-Za-z])', 't', out)
        except Exception:
            pass
        return out

# Optional better PDF extractors for fallback
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except Exception:
    HAS_PDFPLUMBER = False

try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except Exception:
    HAS_PYPDF = False

def _alt_pdf_extract(path: str):
    """Try alternative PDF extraction methods (pdfplumber, pypdf)."""
    docs = []
    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    docs.append(LC_Document(page_content=_fix_text(text), metadata={'source': path, 'page': i+1}))
            if docs:
                return docs
        except Exception:
            pass

    if HAS_PYPDF:
        try:
            reader = PdfReader(path)
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text() or ""
                except Exception:
                    text = ""
                docs.append(LC_Document(page_content=_fix_text(text), metadata={'source': path, 'page': i+1}))
            if docs:
                return docs
        except Exception:
            pass

    return []

# --- 1. LLM Setup (Base Requirement 5) ---
model_name = ("meta-llama/llama-4-scout-17b-16e-instruct")
llm = None
if model_name:
    try:
        llm = ChatGroq(
            model_name=model_name,
            groq_api_key=os.environ.get("GROQ_API_KEY"),
            temperature=0,
        )
        print("Initialized ChatGroq model:", model_name)
    except Exception as e:
        print("Could not initialize ChatGroq:", e)
        llm = None
else:
    print("GROQ_MODEL not set; LLM calls are disabled.")

# --- ENHANCEMENT 1: Sentence-Transformers (Enhancement 1) ---
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

vectorstore = None
doc_store = []
LOG_FILE = Path("query_logs.csv")

def process_pdfs(files):
    """Handles PDF upload, text extraction, and semantic chunking."""
    global vectorstore
    global doc_store
    if not files:
        return "Please upload at least one PDF.", "", ""
    
    all_docs = []
    doc_store = []
    def _resolve_path(f):
        candidates = []
        if isinstance(f, str):
            candidates.append(f)
        if isinstance(f, dict):
            for k in ('tmp_path', 'name', 'file_path', 'path'):
                if k in f:
                    candidates.append(f[k])
        for attr in ('tmp_path', 'name', 'file', 'filename', 'path'):
            if hasattr(f, attr):
                try:
                    candidates.append(getattr(f, attr))
                except Exception:
                    pass

        for c in candidates:
            try:
                if c and Path(c).exists():
                    return str(c)
            except Exception:
                pass

        read = getattr(f, 'read', None)
        if callable(read):
            try:
                data = read()
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=Path(getattr(f, 'name', '')).suffix or '.pdf')
                tmp.write(data)
                tmp.flush()
                tmp.close()
                return tmp.name
            except Exception:
                pass

        return None

    for file in files:
        try:
            path = _resolve_path(file)
            if not path:
                return f"Could not resolve uploaded file path for {getattr(file, 'name', str(file))}.", "", ""

            suffix = Path(path).suffix.lower()

            if suffix == '.pdf':
                loader = PyPDFLoader(path)
                loaded = loader.load()
                raw_has_replacement = any('\ufffd' in (getattr(d, 'page_content', '') or '') for d in loaded)
                cleaned = [LC_Document(page_content=_fix_text(d.page_content), metadata=d.metadata) for d in loaded]

                if raw_has_replacement:
                    alt_docs = _alt_pdf_extract(path)
                    if alt_docs and not any('\ufffd' in (getattr(d, 'page_content', '') or '') for d in alt_docs):
                        cleaned = alt_docs

                all_docs.extend(cleaned)
                joined = "\n\n".join(d.page_content for d in cleaned[:3])
                doc_store.append({'source': path, 'preview': joined, 'pages': len(cleaned)})
            else:
                if HAS_DOCX and suffix in ('.docx', '.doc'):
                    try:
                        doc = docx.Document(path)
                        text = "\n".join(p.text for p in doc.paragraphs)
                    except Exception:
                        raw = Path(path).read_bytes()
                        text = raw.decode('utf-8', errors='replace')
                else:
                    raw = Path(path).read_bytes()
                    text = raw.decode('utf-8', errors='replace')
                
                text = _fix_text(text)
                all_docs.append(LC_Document(page_content=text, metadata={'source': path, 'page': 1}))
                doc_store.append({'source': path, 'preview': text[:1000], 'pages': 1})
        except Exception:
            return f"Error processing {Path(path).name}:\n" + traceback.format_exc(), "", ""
    
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)
    chunks = text_splitter.split_documents(all_docs)
    
    vectorstore = FAISS.from_documents(chunks, embeddings)
    preview_text = "\n\n".join(f"{p['source']}: ({p['pages']} pages)\n{p['preview']}" for p in doc_store)
    combined = "\n\n".join(_fix_text(d.page_content) for d in all_docs)
    try:
        import re
        paragraphs = [p.strip() for p in re.split(r'\n{2,}', combined) if p.strip()]
        summary_lines = []

        def first_n_sentences(text, n=3):
            sents = re.split(r'(?<=[\.\?!])\s+', text.strip())
            return [s for s in sents if s][:n]

        if paragraphs:
            if len(paragraphs) >= 3:
                # take 2 sentences from first 3 paragraphs
                for p in paragraphs[:3]:
                    summary_lines.extend(first_n_sentences(p, 2))
            elif len(paragraphs) == 2:
                for p in paragraphs[:2]:
                    summary_lines.extend(first_n_sentences(p, 3))
            else:
                # single paragraph: try to extract list items first
                single = paragraphs[0]
                items = re.findall(r'(?:\d+\.|•|\-|–)\s*[^\n]+', single)
                if items and len(items) >= 2:
                    summary_lines.extend(items[:10])
                else:
                    summary_lines.extend(first_n_sentences(single, 6))

            # join lines with newlines and enforce length limits
            summary = "\n".join(line.strip() for line in summary_lines if line.strip())
            if not summary:
                summary = " ".join(combined.split()[:200]) + "..."
            if len(summary) > 1200:
                summary = summary[:1200].rstrip() + "..."
        else:
            summary = " ".join(combined.split()[:200]) + "..."
    except Exception:
        summary = " ".join(combined.split()[:200]) + "..."
    return f"✅ Success! Indexed {len(chunks)} chunks from {len(files)} files.", preview_text, summary

def chat_fn(message, history):
    """Retrieves relevant chunks and generates an answer via Groq LLM."""
    global vectorstore
    if vectorstore is None:
        return "⚠️ Please upload and process documents first."

    history_text = ""
    if history:
        start = max(0, len(history) - 6)
        i = start
        parts = []
        while i < len(history):
            msg = history[i]
            role = msg.get('role', '')
            content = msg.get('content', '')
            if role == 'user':
                assistant = ''
                if i + 1 < len(history) and history[i + 1].get('role', '') == 'assistant':
                    assistant = history[i + 1].get('content', '')
                    i += 2
                else:
                    i += 1
                parts.append(f"User: {content}\nAssistant: {assistant}")
            else:
                i += 1
        history_text = "\n".join(parts)

    try:
        top_docs = vectorstore.similarity_search(message, k=3)
        context = "\n\n".join(d.page_content for d in top_docs)
    except Exception:
        top_docs = []
        context = ""

    system_prompt = "You are a helpful assistant. Use the provided context to answer the question."
    user_prompt = f"Context:\n{context}\n\nHistory:\n{history_text}\n\nQuestion:\n{message}"
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    if llm is None:
        if context and context.strip():
            fallback = context.strip()
            if len(fallback) > 1800:
                fallback = fallback[:1800] + "..."
            answer = "[LLM disabled] Local fallback (context summary):\n" + fallback
        else:
            answer = "[LLM disabled] No context available. Set `GROQ_MODEL` and `GROQ_API_KEY`."
    else:
        try:
            response = llm.invoke(messages)
            answer = getattr(response, 'content', None) or str(response)
        except Exception as e:
            try:
                prompt = system_prompt + "\n\n" + user_prompt
                response = llm.invoke(prompt)
                answer = getattr(response, 'content', None) or str(response)
            except Exception as inner_e:
                err_text = str(inner_e)
                if context and context.strip():
                    fallback = context.strip()[:1800]
                    answer = f"[Error: {err_text}]\n\nLocal fallback:\n" + fallback
                else:
                    answer = f"[Error: {err_text}]\nCheck API Key."

    pages = list({str(d.metadata.get('page', '?')) for d in top_docs})
    src_text = f"**Sources:** Page(s) {', '.join(pages)}"
    return f"{answer}\n\n_{src_text}_" if pages else f"{answer}\n\n_(Sources: N/A)_"

def append_log(query, sources, success=True):
    try:
        header = ['timestamp', 'query', 'sources', 'success']
        row = [datetime.datetime.now(datetime.timezone.utc).isoformat(), query, json.dumps(sources), str(success)]
        exists = LOG_FILE.exists()
        with open(LOG_FILE, 'a', newline='', encoding='utf-8') as fh:
            writer = csv.writer(fh)
            if not exists:
                writer.writerow(header)
            writer.writerow(row)
    except Exception:
        pass

def handle_user_message(message, history, enable_logging=False):
    if not message.strip():
        return history, history
        
    if history is None:
        history = []

    answer = chat_fn(message, history)

    history = history + [
        {"role": "user", "content": message},
        {"role": "assistant", "content": answer},
    ]

    try:
        if 'Sources:' in answer:
            src = answer.split('Sources:')[-1].strip()
        else:
            src = ''
    except Exception:
        src = ''
    if enable_logging:
        append_log(message, src, success=True)
    return history, history

def download_history(history):
    if not history:
        return None
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
    i = 0
    while i < len(history):
        user = ''
        assistant = ''
        if history[i].get('role') == 'user':
            user = history[i].get('content', '')
            if i + 1 < len(history) and history[i + 1].get('role') == 'assistant':
                assistant = history[i + 1].get('content', '')
            i += 2
        else:
            if history[i].get('role') == 'assistant':
                assistant = history[i].get('content', '')
            i += 1
        tmp.write(f"User: {user}\nAssistant: {assistant}\n\n")
    tmp.flush()
    tmp.close()
    return tmp.name

# =============================================================================
# PREMIUM PROFESSIONAL GUI CONFIGURATION
# =============================================================================

# Professional Theme with Modern Design System
theme = gr.themes.Soft(
    primary_hue="indigo",
    secondary_hue="blue",
    neutral_hue="slate",
    font=[gr.themes.GoogleFont("Inter"), "SF Pro Display", "ui-sans-serif", "system-ui"],
)

# Premium Custom CSS - Enterprise-Grade Styling
custom_css = """
/* Main Container & Layout */
.gradio-container {
    max-width: 1600px !important;
    margin: 0 auto;
    font-family: 'Inter', sans-serif;
}

/* Premium Header Styling */
.premium-header {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    padding: 2.5rem 2rem;
    border-radius: 16px;
    margin-bottom: 2rem;
    box-shadow: 0 10px 40px rgba(102, 126, 234, 0.2);
}

/* Status Badge Styling */
.status-badge {
    display: inline-block;
    padding: 0.5rem 1rem;
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    color: white;
    border-radius: 20px;
    font-weight: 600;
    font-size: 0.9rem;
    box-shadow: 0 4px 12px rgba(102, 126, 234, 0.3);
}

/* Card Styling */
.premium-card {
    background: white;
    border-radius: 12px;
    padding: 1.5rem;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
    border: 1px solid #e2e8f0;
}

/* Button Enhancements */
button.primary {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
    border: none !important;
    font-weight: 600 !important;
}

/* Input Field Styling */
textarea, input {
    border-radius: 8px !important;
    border: 1.5px solid #e2e8f0 !important;
}

textarea:focus, input:focus {
    border-color: #667eea !important;
    box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1) !important;
}
"""

with gr.Blocks(title="🧠 Premium RAG Intelligence Platform") as demo:
    
    # ===== PREMIUM HEADER =====
    gr.HTML("""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 2.5rem 2rem; border-radius: 16px; margin-bottom: 2rem; box-shadow: 0 10px 40px rgba(102, 126, 234, 0.2);">
            <h1 style="color: white; font-size: 2.5rem; font-weight: 700; margin: 0; text-align: center; text-shadow: 0 2px 4px rgba(0,0,0,0.1);">🧠 Premium RAG Intelligence Platform</h1>
            <p style="color: rgba(255, 255, 255, 0.9); font-size: 1.1rem; margin-top: 0.5rem; text-align: center;">Enterprise-Grade AI-Powered Document Analysis & Knowledge Assistant</p>
        </div>
    """)
    
    # ===== MAIN LAYOUT =====
    with gr.Row():
        # ===== LEFT SIDEBAR =====
        with gr.Column(scale=1):
            gr.Markdown("""
                ### 📁 Knowledge Base Manager
                Upload your documents to build an intelligent knowledge base.
            """)
            
            file_input = gr.File(
                label="📄 Upload Documents",
                file_count="multiple",
                file_types=[".pdf", ".docx", ".doc", ".txt"],
            )
            
            process_btn = gr.Button(
                "🚀 Build Knowledge Base",
                variant="primary",
                size="lg"
            )
            
            gr.Markdown("")
            
            status = gr.Markdown(
                """<div class='status-badge'>⚡ Ready to Process</div>""",
                elem_id="status-msg"
            )
            
            gr.Markdown("---")
            
            with gr.Accordion("⚙️ Advanced Settings", open=False):
                enable_logging = gr.Checkbox(
                    label="📊 Enable Analytics & Logging",
                    value=True,
                    info="Track queries for audit and insights"
                )
                
                gr.Markdown("")
                gr.Markdown("**🛠️ Quick Actions**")
                
                with gr.Row():
                    clear_btn = gr.Button("🗑️ Clear", variant="stop", size="sm")
                    download_btn = gr.Button("💾 Export", variant="secondary", size="sm")
                
                download_file = gr.File(label="Download History", visible=False)
            
            gr.Markdown("""
                ---
                **💡 Pro Tips:**
                - Upload multiple PDFs for comprehensive knowledge
                - Use specific queries for better results
                - Check Document Insights tab for overview
            """)
        
        # ===== RIGHT MAIN CONTENT =====
        with gr.Column(scale=3):
            
            with gr.Tabs():
                # ===== TAB 1: CHAT =====
                with gr.TabItem("💬 AI Assistant", id="chat_tab"):
                    gr.HTML("""
                        <div style="padding: 1rem; background: linear-gradient(to right, #f0f9ff, #e0f2fe); border-radius: 8px; margin-bottom: 1rem;">
                            <p style="margin: 0; color: #0369a1; font-weight: 600;">🤖 Ask anything about your uploaded documents</p>
                        </div>
                    """)
                    
                    chatbot = gr.Chatbot(
                        label="AI Conversation",
                        height=550,
                        avatar_images=(
                            "https://api.dicebear.com/7.x/avataaars/svg?seed=User",
                            "https://api.dicebear.com/7.x/bottts/svg?seed=AI"
                        ),
                    )
                    
                    with gr.Row():
                        message_input = gr.Textbox(
                            placeholder="💭 Type your question here... (e.g., 'Summarize the key findings')",
                            scale=9,
                            show_label=False,
                            container=False
                        )
                        send_btn = gr.Button("➤ Send", variant="primary", scale=1)
                    
                    gr.HTML("""
                        <div style="margin-top: 1rem; padding: 0.75rem; background: #fef3c7; border-left: 4px solid #f59e0b; border-radius: 4px;">
                            <p style="margin: 0; font-size: 0.9rem; color: #92400e;">💡 <strong>Tip:</strong> Be specific in your questions for more accurate answers.</p>
                        </div>
                    """)
                
                # ===== TAB 2: INSIGHTS =====
                with gr.TabItem("📊 Document Intelligence", id="insights_tab"):
                    gr.HTML("""
                        <div style="padding: 1rem; background: linear-gradient(to right, #f0fdf4, #dcfce7); border-radius: 8px; margin-bottom: 1.5rem;">
                            <p style="margin: 0; color: #166534; font-weight: 600;">📈 Comprehensive Document Analysis & Insights</p>
                        </div>
                    """)
                    
                    with gr.Row():
                        with gr.Column():
                            gr.Markdown("""
                                <div style="display: flex; align-items: center; margin-bottom: 1rem;">
                                    <span style="font-size: 1.5rem; margin-right: 0.5rem;">📜</span>
                                    <h3 style="margin: 0;">AI-Generated Summary</h3>
                                </div>
                            """)
                            auto_summary = gr.Textbox(
                                lines=12,
                                interactive=False,
                                placeholder="Upload and process documents to see AI-generated summary...",
                                show_label=False
                            )
                        
                        with gr.Column():
                            gr.Markdown("""
                                <div style="display: flex; align-items: center; margin-bottom: 1rem;">
                                    <span style="font-size: 1.5rem; margin-right: 0.5rem;">🔍</span>
                                    <h3 style="margin: 0;">Raw Content Preview</h3>
                                </div>
                            """)
                            file_preview = gr.Textbox(
                                lines=12,
                                interactive=False,
                                placeholder="Extracted text preview will appear here...",
                                show_label=False
                            )
    
    # ===== FOOTER =====
    gr.HTML("""
        <div style="text-align: center; padding: 2rem 1rem; margin-top: 3rem; border-top: 1px solid #e2e8f0; color: #64748b;">
            <p style="font-size: 0.95rem;"><strong>⚡ Powered by Advanced AI Technology</strong></p>
            <p style="margin-top: 0.5rem; color: #94a3b8;">LangChain • Groq LLM • FAISS Vector Store • Gradio</p>
            <p style="margin-top: 1rem; font-size: 0.85rem; color: #cbd5e1;">© 2025 Premium RAG Platform • Enterprise Edition</p>
        </div>
    """)

    # --- STATE & LOGIC ---
    history_state = gr.State([])

    # Process files and update status markdown
    def process_and_update(files):
        msg, preview, summary = process_pdfs(files)
        return f"System Status: {msg}", preview, summary

    process_btn.click(
        fn=process_and_update, 
        inputs=[file_input], 
        outputs=[status, file_preview, auto_summary]
    )

    # Chat handlers
    message_input.submit(
        fn=handle_user_message, 
        inputs=[message_input, history_state, enable_logging], 
        outputs=[chatbot, history_state]
    ).then(lambda: "", None, message_input)

    send_btn.click(
        fn=handle_user_message, 
        inputs=[message_input, history_state, enable_logging], 
        outputs=[chatbot, history_state]
    ).then(lambda: "", None, message_input)

    # Utility handlers
    download_btn.click(
        fn=download_history, 
        inputs=[history_state], 
        outputs=[download_file]
    ).then(
        lambda: gr.File(visible=True), 
        None, 
        download_file
    )
    clear_btn.click(lambda: ([], "System Status: *Ready*"), None, [chatbot, status])

if __name__ == "__main__":
    # Pass theme and css to launch() for Gradio 6.x compatibility
    demo.launch(
        theme=theme,
        css=custom_css,
        share=False
    )